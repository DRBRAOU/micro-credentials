"""Knowledge retrieval from uploaded documents."""

from pathlib import Path
from typing import Optional, List, Tuple
from docx import Document


def read_docx(file_path: str) -> str:
    """Read text content from a DOCX file."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        return ""


def read_pdf(file_path: str) -> str:
    """Read text content from a PDF file."""
    try:
        import fitz
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception:
        return ""


def read_uploaded_file(uploaded_file) -> str:
    """Read text from an uploaded file (Streamlit UploadedFile)."""
    try:
        if uploaded_file.name.endswith(".txt"):
            return uploaded_file.read().decode("utf-8")
        elif uploaded_file.name.endswith(".docx"):
            import io
            doc = Document(io.BytesIO(uploaded_file.read()))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif uploaded_file.name.endswith(".pdf"):
            try:
                import fitz
                import io
                pdf_bytes = uploaded_file.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except ImportError:
                return "[PDF reading requires PyMuPDF. Install with: pip install pymupdf]"
        else:
            return ""
    except Exception as e:
        return f"[Error reading file: {str(e)}]"


def load_knowledge_context(max_chars: int = 12000) -> str:
    """Load context from ALL knowledge documents (DOCX + PDF).

    Reads BRAOU templates and framework PDFs to provide context for AI generation.

    Args:
        max_chars: Maximum characters to include.

    Returns:
        Combined text from all knowledge documents.
    """
    knowledge_dir = Path(__file__).parent.parent / "knowledge"
    context_parts = []
    total_chars = 0

    # Load ALL supported files: DOCX and PDF
    all_files = sorted(knowledge_dir.glob("*.docx")) + sorted(knowledge_dir.glob("*.pdf"))

    for file_path in all_files:
        if total_chars >= max_chars:
            break

        # Read based on file type
        if file_path.suffix == ".docx":
            content = read_docx(str(file_path))
        elif file_path.suffix == ".pdf":
            content = read_pdf(str(file_path))
        else:
            continue

        if content:
            header = f"\n--- {file_path.stem} ---\n"
            available = max_chars - total_chars - len(header)
            if available > 200:
                truncated = content[:available]
                context_parts.append(header + truncated)
                total_chars += len(header) + len(truncated)

    return "\n".join(context_parts)


def get_loaded_documents_info() -> List[Tuple[str, int]]:
    """Get info about which knowledge documents are loaded and their sizes.

    Returns:
        List of tuples (filename, character_count) for each loaded document.
    """
    knowledge_dir = Path(__file__).parent.parent / "knowledge"
    results = []

    all_files = sorted(knowledge_dir.glob("*.docx")) + sorted(knowledge_dir.glob("*.pdf"))

    for file_path in all_files:
        if file_path.suffix == ".docx":
            content = read_docx(str(file_path))
        elif file_path.suffix == ".pdf":
            content = read_pdf(str(file_path))
        else:
            continue

        results.append((file_path.name, len(content)))

    return results
